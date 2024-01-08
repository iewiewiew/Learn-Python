# !/usr/bin/env python
# -*- coding:utf-8 -*-

"""
@author       weimenghua
@time         2023/9/4 18:14
@description  docx 文件操作
"""

import random
import time

from docx import Document


def docx_ops(path):
    # 创建一个新的文档对象
    doc = Document()

    # 添加标题
    doc.add_heading('示例文档', level=1)

    # 添加段落
    doc.add_paragraph('这是一个示例段落。')

    # 添加有序列表
    doc.add_paragraph('有序列表:', style='List Number')
    list_items = ['第一项', '第二项', '第三项']
    for item in list_items:
        doc.add_paragraph(item, style='List Number 2')

    # 添加无序列表
    doc.add_paragraph('无序列表:', style='List Bullet')
    list_items = ['项目 A', '项目 B', '项目 C']
    for item in list_items:
        doc.add_paragraph(item, style='List Bullet 2')
    doc.add_paragraph("{:05d}".format(random.Random().randrange(10000, 100000)))
    doc.add_paragraph(time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(time.time())))

    # 保存文档
    doc.save(path)


if __name__ == '__main__':
    path = 'example.docx'
    docx_ops(path)